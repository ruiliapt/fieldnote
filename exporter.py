"""
导出模块 - 负责将语料导出为各种格式
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict


class TextFormatter:
    """文本格式化类 - 生成对齐的语言学格式文本"""
    
    @staticmethod
    def calculate_display_width(text: str) -> int:
        """
        计算文本的显示宽度（考虑等宽字体）
        在等宽字体中：
        - 中文、日文等宽字符：2
        - 英文、数字、标点：1
        - 上标数字（音调）：0.5（计为1，但多个上标算一起）
        """
        import unicodedata
        
        width = 0
        i = 0
        text_len = len(text)
        
        while i < text_len:
            char = text[i]
            code = ord(char)
            
            # 中日韩统一表意文字
            if 0x4E00 <= code <= 0x9FFF:
                width += 2
            # 中日韩扩展
            elif 0x3400 <= code <= 0x4DBF:
                width += 2
            # 全角字符
            elif 0xFF00 <= code <= 0xFFEF:
                width += 2
            # 上标数字（音调标记）- 统一计为0.5宽度
            elif 0x2070 <= code <= 0x209F or 0x00B2 <= code <= 0x00B3:
                # 连续的上标数字一起计算
                superscript_count = 0
                while i < text_len and (0x2070 <= ord(text[i]) <= 0x209F or 0x00B2 <= ord(text[i]) <= 0x00B3):
                    superscript_count += 1
                    i += 1
                i -= 1  # 回退一个，因为循环会+1
                # 每2个上标字符算1个宽度
                width += (superscript_count + 1) // 2
            # 组合变音符号（紧跟在字母后面的）
            elif unicodedata.category(char) in ['Mn', 'Mc', 'Me']:
                width += 0
            # 其他字符
            else:
                width += 1
            
            i += 1
        
        return width
    
    @staticmethod
    def align_words(source_words: List[str], gloss_words: List[str]) -> tuple:
        """
        对齐原文和词汇分解（紧凑型完美左对齐）
        
        Args:
            source_words: 原文词列表
            gloss_words: 词汇分解列表
            
        Returns:
            (对齐后的原文行, 对齐后的gloss行)
        """
        # 确保两个列表长度相同
        max_len = max(len(source_words), len(gloss_words))
        source_words = source_words + [''] * (max_len - len(source_words))
        gloss_words = gloss_words + [''] * (max_len - len(gloss_words))
        
        # 计算每个位置需要的实际宽度（紧凑模式）
        column_widths = []
        for src, gls in zip(source_words, gloss_words):
            # 计算字符长度（包括上标）
            src_len = len(src)
            gls_len = len(gls)
            
            # 取较大者，并加上最少2个空格作为间距
            col_width = max(src_len, gls_len) + 2
            column_widths.append(col_width)
        
        # 对齐每一列
        aligned_source_parts = []
        aligned_gloss_parts = []
        
        for src, gls, col_width in zip(source_words, gloss_words, column_widths):
            # 左对齐，右侧填充空格到列宽
            src_padded = src.ljust(col_width)
            gls_padded = gls.ljust(col_width)
            
            aligned_source_parts.append(src_padded)
            aligned_gloss_parts.append(gls_padded)
        
        # 连接所有列
        return (''.join(aligned_source_parts), ''.join(aligned_gloss_parts))
    
    @staticmethod
    def format_entry(entry: Dict, show_numbering: bool = True, number_format: str = "()",
                    max_words_per_line: int = 10, include_chinese: bool = False) -> str:
        """
        格式化单条语料为标准语言学格式（Leipzig Glossing Rules）
        支持长句自动分行
        
        Args:
            entry: 语料记录
            show_numbering: 是否显示编号
            number_format: 编号格式 "()" 或 "."
            max_words_per_line: 每行最多显示的词数（默认10）
            include_chinese: 是否包含汉字注释字段
            
        Returns:
            格式化后的文本
        """
        lines = []
        
        # 准备编号
        numbering_text = ""
        if show_numbering:
            example_id = entry.get('example_id', '')
            if example_id:
                if number_format == "()":
                    numbering_text = f"({example_id}) "
                else:
                    numbering_text = f"{example_id}. "
            else:
                entry_id = entry.get('id', '')
                if number_format == "()":
                    numbering_text = f"({entry_id}) "
                else:
                    numbering_text = f"{entry_id}. "
        
        # 分词（按空格分割）
        source_text = entry.get('source_text', '').strip()
        gloss = entry.get('gloss', '').strip()
        translation = entry.get('translation', '').strip()
        notes = entry.get('notes', '').strip()
        
        # 汉字字段（如果包含）
        source_text_cn = entry.get('source_text_cn', '').strip() if include_chinese else ""
        gloss_cn = entry.get('gloss_cn', '').strip() if include_chinese else ""
        translation_cn = entry.get('translation_cn', '').strip() if include_chinese else ""
        
        # 计算缩进空格数（编号的长度）
        indent_spaces = ' ' * len(numbering_text)
        
        # 原文行和注释行 - 进行词对齐（支持自动分行）
        source_has_words = len(source_text.split()) > 1
        gloss_has_words = len(gloss.split()) > 1
        
        if source_has_words and gloss_has_words:
            # 多词情况：编号和原文在同一行
            source_words = source_text.split()
            gloss_words = gloss.split()
            
            # 如果原文最后一个词是单独的句号，将其合并到前一个词
            if len(source_words) > 1 and source_words[-1] == '.':
                source_words[-2] = source_words[-2] + '.'
                source_words.pop()
            
            # 确保词数相同
            max_len = max(len(source_words), len(gloss_words))
            source_words = source_words + [''] * (max_len - len(source_words))
            gloss_words = gloss_words + [''] * (max_len - len(gloss_words))
            
            # 如果词数超过每行最大值，分行显示
            if len(source_words) > max_words_per_line:
                # 计算缩进
                if numbering_text:
                    indent = ' ' * len(numbering_text)
                else:
                    indent = '    '
                
                # 分批处理
                for i in range(0, len(source_words), max_words_per_line):
                    batch_source = source_words[i:i + max_words_per_line]
                    batch_gloss = gloss_words[i:i + max_words_per_line]
                    
                    aligned_source, aligned_gloss = TextFormatter.align_words(batch_source, batch_gloss)
                    # 第一批添加编号
                    if i == 0 and numbering_text:
                        lines.append(f"{numbering_text}{aligned_source}")
                    else:
                        lines.append(f"{indent}{aligned_source}")
                    lines.append(f"{indent}{aligned_gloss}")
                    
                    # 如果还有下一批，添加空行分隔
                    if i + max_words_per_line < len(source_words):
                        lines.append("")
            else:
                # 一行能显示完
                aligned_source, aligned_gloss = TextFormatter.align_words(source_words, gloss_words)
                # 编号和原文在同一行
                if numbering_text:
                    lines.append(f"{numbering_text}{aligned_source}")
                    # 计算缩进（编号的长度）
                    indent = ' ' * len(numbering_text)
                else:
                    lines.append(f"    {aligned_source}")
                    indent = '    '
                
                # 原文(汉字) 紧跟原文
                if source_text_cn:
                    lines.append(f"{indent}【原文(汉字)】{source_text_cn}")
                
                lines.append(f"{indent}{aligned_gloss}")
                # 词汇分解(汉字) 紧跟词汇分解
                if gloss_cn:
                    lines.append(f"{indent}【词汇分解(汉字)】{gloss_cn}")
            
            # 翻译行（使用编号的缩进）
            if numbering_text:
                indent = ' ' * len(numbering_text)
            else:
                indent = '    '
            
            if translation:
                if notes:
                    lines.append(f"{indent}'{translation}' ({notes})")
                else:
                    lines.append(f"{indent}'{translation}'")
            elif notes:
                lines.append(f"{indent}({notes})")
            
            # 翻译(汉字) 紧跟翻译
            if translation_cn:
                lines.append(f"{indent}【翻译(汉字)】{translation_cn}")
        else:
            # 单词情况：编号和原文在同一行
            lines.append(f"{numbering_text}{source_text}")
            # 原文(汉字) 紧跟原文
            if source_text_cn:
                lines.append(f"{indent_spaces}【原文(汉字)】{source_text_cn}")
            
            # gloss 行（缩进对齐）
            if gloss:
                lines.append(f"{indent_spaces}{gloss}")
            # 词汇分解(汉字) 紧跟词汇分解
            if gloss_cn:
                lines.append(f"{indent_spaces}【词汇分解(汉字)】{gloss_cn}")
            
            # 翻译行（缩进对齐）
            if translation:
                if notes:
                    lines.append(f"{indent_spaces}'{translation}' ({notes})")
                else:
                    lines.append(f"{indent_spaces}'{translation}'")
            elif notes:
                lines.append(f"{indent_spaces}({notes})")
            # 翻译(汉字) 紧跟翻译
            if translation_cn:
                lines.append(f"{indent_spaces}【翻译(汉字)】{translation_cn}")
        
        return '\n'.join(lines)
    
    @staticmethod
    def format_entries(entries: List[Dict], show_numbering: bool = True, 
                      number_format: str = "()", include_chinese: bool = False) -> str:
        """
        格式化多条语料
        
        Args:
            entries: 语料记录列表
            show_numbering: 是否显示编号
            number_format: 编号格式 "()" 或 "."
            include_chinese: 是否包含汉字注释字段
            
        Returns:
            格式化后的文本
        """
        formatted = []
        for entry in entries:
            formatted.append(TextFormatter.format_entry(entry, show_numbering, number_format, 
                                                       include_chinese=include_chinese))
            formatted.append('')  # 空行分隔每条语料
        
        return '\n'.join(formatted)


class WordExporter:
    """Word文档导出类"""
    
    def __init__(self):
        """初始化导出器"""
        self.doc = Document()
    
    def _set_cell_properties(self, cell, font_size, is_content_cell=True, font_name=None):
        """设置单元格的XML属性和段落格式
        
        Args:
            cell: 单元格对象
            font_size: 字体大小
            is_content_cell: 是否是内容单元格（有文字的），空单元格不设置字体
            font_name: 字体名称（可选）
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '0')
        tcW.set(qn('w:type'), 'auto')
        tcPr.append(tcW)
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'top')
        tcPr.append(vAlign)
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom', 'left', 'right']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if is_content_cell:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if font_name and font_name != "系统默认":
                        run.font.name = font_name
    
    def _split_words_by_width(self, words: List[str], max_width: int = 80) -> List[List[str]]:
        """根据显示宽度自动将词分行
        
        Args:
            words: 词列表
            max_width: 每行最大显示宽度（等宽字体字符数）
            
        Returns:
            分行后的词列表，每个子列表代表一行
        """
        if not words:
            return [[]]
        
        lines = []
        current_line = []
        current_width = 0
        
        for word in words:
            word_width = TextFormatter.calculate_display_width(word)
            # 每个词之间留一个空格的宽度（约2个字符）
            space_width = 2 if current_line else 0
            
            if current_width + space_width + word_width > max_width and current_line:
                # 当前行放不下了，开始新行
                lines.append(current_line)
                current_line = [word]
                current_width = word_width
            else:
                # 当前行还能放下
                current_line.append(word)
                current_width += space_width + word_width
        
        if current_line:
            lines.append(current_line)
        
        return lines if lines else [[]]
    
    def export(self, entries: List[Dict], output_path: str, 
               table_width: float = 5.0,
               font_size: int = 10,
               line_spacing: float = 1.15,
               show_numbering: bool = True,
               entries_per_page: int = 10,
               include_chinese: bool = False,
               max_words_per_line: int = 10,
               font_config: Dict = None) -> bool:
        """
        导出语料到Word文档（词对词对齐，透明表格）
        
        Args:
            entries: 语料记录列表
            output_path: 输出文件路径
            table_width: 表格宽度（英寸）
            font_size: 字体大小（磅）
            line_spacing: 行距
            show_numbering: 是否显示编号
            entries_per_page: 每页语料数（暂未实现分页）
            include_chinese: 是否包含汉字注释字段
            font_config: 字体配置字典，包含各字段的字体名称和大小
            
        Returns:
            是否导出成功
        """
        try:
            # 初始化字体配置
            if font_config is None:
                font_config = {}
            
            # 设置默认字体配置
            source_font = font_config.get("source_text", "Doulos SIL Compact")
            source_size = font_config.get("source_text_size", 12)
            gloss_font = font_config.get("gloss", "Charis SIL Compact")
            gloss_size = font_config.get("gloss_size", 11)
            translation_font = font_config.get("translation", None)
            translation_size = font_config.get("translation_size", 11)
            chinese_font = font_config.get("chinese", None)
            chinese_size = font_config.get("chinese_size", 10)
            
            self.doc = Document()
            
            for idx, entry in enumerate(entries, 1):
                # 准备编号（如果需要）
                numbering_text = ""
                if show_numbering:
                    example_id = entry.get('example_id', '')
                    if example_id:
                        numbering_text = f"({example_id}) "
                    else:
                        numbering_text = f"({idx}) "
                
                # 分词
                source_text = entry.get('source_text', '').strip()
                gloss = entry.get('gloss', '').strip()
                translation = entry.get('translation', '').strip()
                notes = entry.get('notes', '').strip()
                
                # 汉字字段（如果包含）
                source_text_cn = entry.get('source_text_cn', '').strip() if include_chinese else ""
                gloss_cn = entry.get('gloss_cn', '').strip() if include_chinese else ""
                translation_cn = entry.get('translation_cn', '').strip() if include_chinese else ""
                
                # 如果有多个词（分词后长度>1），进行词对齐
                source_has_words = ' ' in source_text and len(source_text.split()) > 1
                gloss_has_words = ' ' in gloss and len(gloss.split()) > 1
                if source_has_words and gloss_has_words:
                    source_words = source_text.split()
                    gloss_words = gloss.split()
                    
                    # 如果原文最后一个词是单独的句号，将其合并到前一个词
                    if len(source_words) > 1 and source_words[-1] == '.':
                        source_words[-2] = source_words[-2] + '.'
                        source_words.pop()
                    
                    # 分词处理汉字字段
                    source_words_cn = source_text_cn.split() if (include_chinese and source_text_cn) else []
                    gloss_words_cn = gloss_cn.split() if (include_chinese and gloss_cn) else []
                    
                    # 根据显示宽度自动分行
                    max_line_width = 80  # 每行最大显示宽度（等宽字体字符数）
                    source_lines_list = self._split_words_by_width(source_words, max_line_width)
                    gloss_lines_list = self._split_words_by_width(gloss_words, max_line_width)
                    
                    # 汉字字段也进行分行（和原文/gloss保持相同的分行方式）
                    source_cn_lines_list = self._split_words_by_width(source_words_cn, max_line_width) if source_words_cn else []
                    gloss_cn_lines_list = self._split_words_by_width(gloss_words_cn, max_line_width) if gloss_words_cn else []
                    
                    # 计算表格尺寸
                    source_line_count = len(source_lines_list)
                    gloss_line_count = len(gloss_lines_list)
                    
                    # 计算总行数：原文 + 原文(汉字) + gloss + gloss(汉字) + 翻译 + 翻译(汉字)
                    total_rows = source_line_count + gloss_line_count + 1  # 原文 + gloss + 翻译
                    if include_chinese:
                        if source_words_cn:
                            total_rows += len(source_cn_lines_list)  # 原文(汉字)行数与原文相同
                        if gloss_words_cn:
                            total_rows += len(gloss_cn_lines_list)  # 词汇分解(汉字)行数与gloss相同
                        if translation_cn:
                            total_rows += 1  # 翻译(汉字)单行合并
                    
                    # 找出所有行中最长的一行（词数最多）
                    all_line_lists = [source_lines_list, gloss_lines_list]
                    if source_cn_lines_list:
                        all_line_lists.append(source_cn_lines_list)
                    if gloss_cn_lines_list:
                        all_line_lists.append(gloss_cn_lines_list)
                    
                    max_cols_in_line = max(
                        max(len(line) for line in line_list)
                        for line_list in all_line_lists
                    )
                    table_cols = max_cols_in_line + 1  # +1 for numbering column
                    
                    table = self.doc.add_table(rows=total_rows, cols=table_cols)
                    
                    # 设置表格属性
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    # 设置无边框
                    tblBorders = OxmlElement('w:tblBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'none')
                        border.set(qn('w:sz'), '0')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), 'auto')
                        tblBorders.append(border)
                    tblPr.append(tblBorders)
                    
                    # 表格布局：自适应
                    tblLayout = OxmlElement('w:tblLayout')
                    tblLayout.set(qn('w:type'), 'auto')
                    tblPr.append(tblLayout)
                    
                    # 表格宽度：自适应
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '0')
                    tblW.set(qn('w:type'), 'auto')
                    tblPr.append(tblW)
                    
                    # 单元格间距：0
                    tblCellSpacing = OxmlElement('w:tblCellSpacing')
                    tblCellSpacing.set(qn('w:w'), '0')
                    tblCellSpacing.set(qn('w:type'), 'dxa')
                    tblPr.append(tblCellSpacing)
                    
                    # 单元格边距
                    tblCellMar = OxmlElement('w:tblCellMar')
                    for side in ['left', 'right']:
                        margin = OxmlElement(f'w:{side}')
                        margin.set(qn('w:w'), '50')
                        margin.set(qn('w:type'), 'dxa')
                        tblCellMar.append(margin)
                    for side in ['top', 'bottom']:
                        margin = OxmlElement(f'w:{side}')
                        margin.set(qn('w:w'), '0')
                        margin.set(qn('w:type'), 'dxa')
                        tblCellMar.append(margin)
                    tblPr.append(tblCellMar)
                    
                    # 设置所有行自动调整高度
                    for row in table.rows:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        trHeight = trPr.find(qn('w:trHeight'))
                        if trHeight is not None:
                            trPr.remove(trHeight)
                        trHeight = OxmlElement('w:trHeight')
                        trHeight.set(qn('w:hRule'), 'auto')
                        trPr.append(trHeight)
                    
                    # 关键：清空所有行的所有单元格
                    for row_idx in range(total_rows):
                        for cell in table.rows[row_idx].cells:
                            cell._element.clear_content()
                            cell.text = ''  # 添加空段落
                    
                    # 追踪当前填充到哪一行
                    current_row = 0
                    
                    # 填充原文（多行，自动换行）
                    for line_idx, line_words in enumerate(source_lines_list):
                        # 第0列：第一行放编号，其他行留空
                        cell = table.rows[current_row].cells[0]
                        if line_idx == 0 and numbering_text:
                            cell.text = numbering_text
                            self._set_cell_properties(cell, source_size, is_content_cell=True, font_name=source_font)
                        else:
                            cell.text = ''
                            self._set_cell_properties(cell, source_size, is_content_cell=False)
                        
                        # 第1到N列：填充这一行的词
                        for col_idx, word in enumerate(line_words):
                            cell = table.rows[current_row].cells[col_idx + 1]
                            cell.text = word
                            self._set_cell_properties(cell, source_size, is_content_cell=True, font_name=source_font)
                        
                        current_row += 1
                    
                    # 填充原文(汉字)（多行，自动换行，每个词一个单元格）
                    for line_idx, line_words in enumerate(source_cn_lines_list):
                        # 第0列：留空
                        cell = table.rows[current_row].cells[0]
                        cell.text = ''
                        self._set_cell_properties(cell, chinese_size, is_content_cell=False)
                        
                        # 第1到N列：填充这一行的汉字词
                        for col_idx, word in enumerate(line_words):
                            cell = table.rows[current_row].cells[col_idx + 1]
                            cell.text = word
                            self._set_cell_properties(cell, chinese_size, is_content_cell=True, font_name=chinese_font)
                        
                        current_row += 1
                    
                    # 填充gloss（多行，自动换行）
                    for line_idx, line_words in enumerate(gloss_lines_list):
                        # 第0列：留空
                        cell = table.rows[current_row].cells[0]
                        cell.text = ''
                        self._set_cell_properties(cell, gloss_size, is_content_cell=False)
                        
                        # 第1到N列：填充这一行的gloss词
                        for col_idx, word in enumerate(line_words):
                            cell = table.rows[current_row].cells[col_idx + 1]
                            cell.text = word
                            self._set_cell_properties(cell, gloss_size, is_content_cell=True, font_name=gloss_font)
                        
                        current_row += 1
                    
                    # 填充词汇分解(汉字)（多行，自动换行，每个词一个单元格）
                    for line_idx, line_words in enumerate(gloss_cn_lines_list):
                        # 第0列：留空
                        cell = table.rows[current_row].cells[0]
                        cell.text = ''
                        self._set_cell_properties(cell, chinese_size, is_content_cell=False)
                        
                        # 第1到N列：填充这一行的汉字词
                        for col_idx, word in enumerate(line_words):
                            cell = table.rows[current_row].cells[col_idx + 1]
                            cell.text = word
                            self._set_cell_properties(cell, chinese_size, is_content_cell=True, font_name=chinese_font)
                        
                        current_row += 1
                    
                    # 填充翻译行（合并所有列）
                    # 第0列：留空
                    cell = table.rows[current_row].cells[0]
                    cell.text = ''
                    self._set_cell_properties(cell, translation_size, is_content_cell=False)
                    
                    if translation:
                        # 合并第1列到最后一列
                        merged_cell = table.rows[current_row].cells[1]
                        for col_idx in range(2, table_cols):
                            merged_cell.merge(table.rows[current_row].cells[col_idx])
                        merged_cell.text = f"'{translation}'"
                        self._set_cell_properties(merged_cell, translation_size, is_content_cell=True, font_name=translation_font)
                    
                    current_row += 1
                    
                    # 填充翻译(汉字)行（如果有）
                    if include_chinese and translation_cn:
                        # 第0列：留空
                        cell = table.rows[current_row].cells[0]
                        cell.text = ''
                        self._set_cell_properties(cell, chinese_size, is_content_cell=False)
                        
                        # 合并第1列到最后一列
                        merged_cell = table.rows[current_row].cells[1]
                        for col_idx in range(2, table_cols):
                            merged_cell.merge(table.rows[current_row].cells[col_idx])
                        merged_cell.text = f"'{translation_cn}'"
                        self._set_cell_properties(merged_cell, chinese_size, is_content_cell=True, font_name=chinese_font)
                    
                    # 汉字字段已经集成到表格中
                    
                else:
                    # 如果没有分词，使用段落格式（不用表格）
                    from docx.shared import Cm
                    
                    # 计算缩进量（编号的宽度）
                    # 假设每个字符约 0.15cm，根据编号长度计算
                    indent_cm = len(numbering_text) * 0.15 if numbering_text else 0
                    
                    # 第一行：编号 + 原文
                    p1 = self.doc.add_paragraph()
                    if numbering_text:
                        p1.add_run(numbering_text).font.size = Pt(font_size)
                    p1.add_run(source_text).font.size = Pt(font_size)
                    p1.paragraph_format.space_before = Pt(0)
                    p1.paragraph_format.space_after = Pt(0)
                    p1.paragraph_format.line_spacing = 1.0
                    
                    # 原文(汉字) 紧跟原文（只在include_chinese时显示）
                    if include_chinese and source_text_cn:
                        p_cn = self.doc.add_paragraph(source_text_cn)
                        p_cn.runs[0].font.size = Pt(font_size - 1)
                        p_cn.paragraph_format.left_indent = Cm(indent_cm)
                        p_cn.paragraph_format.space_before = Pt(0)
                        p_cn.paragraph_format.space_after = Pt(0)
                        p_cn.paragraph_format.line_spacing = 1.0
                    
                    # 第二行：gloss（左缩进与原文对齐）
                    if gloss:
                        p2 = self.doc.add_paragraph(gloss)
                        p2.runs[0].font.size = Pt(font_size)
                        p2.paragraph_format.left_indent = Cm(indent_cm)
                        p2.paragraph_format.space_before = Pt(0)
                        p2.paragraph_format.space_after = Pt(0)
                        p2.paragraph_format.line_spacing = 1.0
                    
                    # 词汇分解(汉字) 紧跟词汇分解（只在include_chinese时显示）
                    if include_chinese and gloss_cn:
                        p_cn = self.doc.add_paragraph(gloss_cn)
                        p_cn.runs[0].font.size = Pt(font_size - 1)
                        p_cn.paragraph_format.left_indent = Cm(indent_cm)
                        p_cn.paragraph_format.space_before = Pt(0)
                        p_cn.paragraph_format.space_after = Pt(0)
                        p_cn.paragraph_format.line_spacing = 1.0
                    
                    # 第三行：翻译（左缩进与原文对齐）
                    if translation:
                        p3 = self.doc.add_paragraph(f"'{translation}'")
                        p3.runs[0].font.size = Pt(font_size)
                        p3.paragraph_format.left_indent = Cm(indent_cm)
                        p3.paragraph_format.space_before = Pt(0)
                        p3.paragraph_format.space_after = Pt(0)
                        p3.paragraph_format.line_spacing = 1.0
                    
                    # 翻译(汉字) 紧跟翻译（只在include_chinese时显示）
                    if include_chinese and translation_cn:
                        p_cn = self.doc.add_paragraph(f"'{translation_cn}'")
                        p_cn.runs[0].font.size = Pt(font_size - 1)
                        p_cn.paragraph_format.left_indent = Cm(indent_cm)
                        p_cn.paragraph_format.space_before = Pt(0)
                        p_cn.paragraph_format.space_after = Pt(0)
                        p_cn.paragraph_format.line_spacing = 1.0
                
                # 添加备注（在表格外，与原文第一个词对齐）
                if notes:
                    note_p = self.doc.add_paragraph(f"    ({notes})")
                    note_p.runs[0].font.size = Pt(font_size - 1)
                    note_p.runs[0].italic = True
                    note_p.paragraph_format.space_before = Pt(0)
                    note_p.paragraph_format.space_after = Pt(0)
                    note_p.paragraph_format.line_spacing = 1.0
                
                # 添加段落间距
                self.doc.add_paragraph()
            
            # 保存文档
            self.doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"导出失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def export_custom_format(self, entries: List[Dict], output_path: str,
                            format_template: str = None) -> bool:
        """
        使用自定义格式导出
        
        Args:
            entries: 语料记录列表
            output_path: 输出文件路径
            format_template: 自定义格式模板（未来功能）
            
        Returns:
            是否导出成功
        """
        # 未来可扩展的自定义格式导出
        return self.export(entries, output_path)

